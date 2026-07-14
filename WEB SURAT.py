import streamlit as st
import pandas as pd
from docx import Document
import io
import zipfile

st.set_page_config(
    page_title="Generator Surat Otomatis",
    page_icon="✉️",
    layout="centered"
)

st.title("✉️ Generator Surat Otomatis")
st.write("Aplikasi internal tim untuk membuat surat massal secara instan dari Excel dan Word.")

# Bagian Panduan Singkat untuk Tim
with st.expander("ℹ️ Cara Menggunakan Aplikasi Ini (Klik untuk Membuka)"):
    st.markdown("""
    1. **Siapkan Template Word (`.docx`):** Pastikan di dalam dokumen surat Anda sudah ada penanda seperti `{{Nama}}`, `{{Alamat}}`, atau `{{No_Surat}}`.
    2. **Siapkan Data Excel (`.xlsx`):** Pastikan baris pertama/header Excel Anda memiliki nama kolom yang **sama persis** dengan penanda di Word (misal kolom: *Nama*, *Alamat*, *No_Surat*).
    3. **Unggah Kedua File:** Masukkan file ke kolom unggah di bawah ini.
    4. **Download ZIP:** Klik tombol generate, lalu unduh folder ZIP berisi semua surat yang sudah jadi.
    """)

# Kolom Unggah File
st.subheader("1. Unggah Berkas")
col1, col2 = st.columns(2)

with col1:
    excel_file = st.file_uploader("Unggah File Excel (Data)", type=["xlsx", "xls"])
with col2:
    word_file = st.file_uploader("Unggah Template Word (Docx)", type=["docx"])

def replace_placeholders(doc, data_row):
    # Ganti teks di paragraf
    for paragraph in doc.paragraphs:
        for key, value in data_row.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))

    # Ganti teks di tabel
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data_row.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in paragraph.text:
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(value))
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, str(value))

if excel_file and word_file:
    try:
        df = pd.read_excel(excel_file)
        
        st.subheader("2. Pengaturan Dokumen")
        st.success("File Excel berhasil dibaca!")
        
        # Pilihan nama file hasil
        columns = df.columns.tolist()
        naming_col = st.selectbox(
            "Pilih kolom Excel yang ingin dijadikan Nama File Surat (misal: Nama):",
            options=columns
        )
        
        if st.button("Proses & Generate Surat", type="primary"):
            zip_buffer = io.BytesIO()
            
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                template_bytes = word_file.read()
                
                progress_bar = st.progress(0)
                status_text = st.empty()
                total_rows = len(df)
                
                for index, row in df.iterrows():
                    progress = (index + 1) / total_rows
                    progress_bar.progress(progress)
                    status_text.text(f"Memproses surat ke-{index+1}...")
                    
                    data_row = row.to_dict()
                    doc = Document(io.BytesIO(template_bytes))
                    
                    replace_placeholders(doc, data_row)
                    
                    out_io = io.BytesIO()
                    doc.save(out_io)
                    out_io.seek(0)
                    
                    clean_filename = str(row[naming_col]).replace("/", "_").replace("\\", "_")
                    filename = f"Surat_{clean_filename}.docx"
                    
                    zip_file.writestr(filename, out_io.getvalue())
            
            zip_buffer.seek(0)
            st.success("🎉 Semua surat selesai dibuat!")
            
            st.download_button(
                label="📥 Download Semua Surat (ZIP)",
                data=zip_buffer,
                file_name="Surat_Otomatis_Tim.zip",
                mime="application/zip"
            )
            
    except Exception as e:
        st.error(f"Terjadi kesalahan: {e}")
else:
    st.info("Menunggu Anda mengunggah file Excel dan Word...")